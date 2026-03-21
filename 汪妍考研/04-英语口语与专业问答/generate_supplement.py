#!/usr/bin/env python3
"""生成护理复试材料补充包 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# 字体配置：中文宋体，英文 Times New Roman，标题用黑体
CN_BODY = '宋体'
CN_HEADING = '黑体'
EN_FONT = 'Times New Roman'

# ── helpers ──────────────────────────────────────────────

def set_run_font(run, cn_font=CN_BODY, en_font=EN_FONT):
    """统一设置 run 的中英文字体"""
    run.font.name = en_font
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), cn_font)

def init_doc_styles(doc):
    """初始化文档默认样式的字体"""
    # Normal 样式
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), CN_BODY)
    # Heading 样式
    for i in range(1, 5):
        key = f'Heading {i}'
        if key in doc.styles:
            hs = doc.styles[key]
            hs.font.name = EN_FONT
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), CN_HEADING)

def set_cell_shading(cell, color_hex):
    """给单元格设置背景色"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        set_run_font(run, cn_font=CN_HEADING)
    return h

def add_para(doc, text, bold=False, size=10.5, space_after=6, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    set_run_font(run)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                set_run_font(run, cn_font=CN_HEADING)
        set_cell_shading(cell, 'D6E4F0')
    # data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    set_run_font(run)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ══════════════════════════════════════════════════════════
#  PART 1 : 难词标注速查表
# ══════════════════════════════════════════════════════════

VOCAB_DATA = {
    "练习1 AI在护理中的整合": [
        ("encompasses", "/ɪnˈkʌmpəsɪz/", "涵盖"),
        ("cognitive", "/ˈkɒɡnɪtɪv/", "认知的"),
        ("replication", "/ˌreplɪˈkeɪʃn/", "复制"),
        ("robotics", "/rəʊˈbɒtɪks/", "机器人技术"),
        ("escalating", "/ˈeskəleɪtɪŋ/", "不断升级的"),
        ("impediments", "/ɪmˈpedɪmənts/", "障碍/阻碍"),
        ("surpasses", "/sərˈpæsɪz/", "超越"),
        ("regulatory", "/ˈreɡjələtɔːri/", "监管的"),
    ],
    "练习2 ChatGPT在护理中的应用": [
        ("constraints", "/kənˈstreɪnts/", "约束/限制"),
        ("simultaneously", "/ˌsaɪməlˈteɪniəsli/", "同时地 ★难读"),
        ("intensifying", "/ɪnˈtensɪfaɪɪŋ/", "加剧"),
        ("comprehension", "/ˌkɒmprɪˈhenʃn/", "理解力"),
        ("initiatives", "/ɪˈnɪʃətɪvz/", "举措/倡议"),
        ("projections", "/prəˈdʒekʃnz/", "预测"),
    ],
    "练习3 护理信息学在AI时代的演进": [
        ("informatics", "/ˌɪnfərˈmætɪks/", "信息学 ★常考"),
        ("leverages", "/ˈlevərɪdʒɪz/", "利用/借助"),
        ("personalized", "/ˈpɜːsənəlaɪzd/", "个性化的"),
        ("allocation", "/ˌæləˈkeɪʃn/", "分配"),
        ("advancements", "/ədˈvɑːnsmənts/", "进展"),
        ("transforming", "/trænsˈfɔːmɪŋ/", "变革"),
    ],
    "练习4 人文护理管理策略": [
        ("humanized", "/ˈhjuːmənaɪzd/", "人性化的"),
        ("explicitly", "/ɪkˈsplɪsɪtli/", "明确地"),
        ("prioritizes", "/praɪˈɒrɪtaɪzɪz/", "优先考虑"),
        ("contradiction", "/ˌkɒntrəˈdɪkʃn/", "矛盾"),
        ("quantitative", "/ˈkwɒntɪtətɪv/", "定量的 ★难读"),
        ("overlook", "/ˌəʊvəˈlʊk/", "忽视"),
    ],
    "练习5 同情疲劳与人文关怀能力": [
        ("compassion", "/kəmˈpæʃn/", "同情/怜悯"),
        ("fatigue", "/fəˈtiːɡ/", "疲劳"),
        ("therapeutic", "/ˌθerəˈpjuːtɪk/", "治疗的 ★θ发音"),
        ("pioneering", "/ˌpaɪəˈnɪərɪŋ/", "开创性的"),
        ("empathy", "/ˈempəθi/", "共情 ★常考"),
        ("competency", "/ˈkɒmpɪtənsi/", "胜任力"),
        ("deliberate", "/dɪˈlɪbərət/", "有意识的"),
    ],
    "练习6 护士主导的姑息护理": [
        ("diminished", "/dɪˈmɪnɪʃt/", "减少的"),
        ("resilience", "/rɪˈzɪliəns/", "韧性/复原力 ★高频"),
        ("compromised", "/ˈkɒmprəmaɪzd/", "受损的"),
        ("incidence", "/ˈɪnsɪdəns/", "发病率"),
        ("alleviating", "/əˈliːvieɪtɪŋ/", "缓解 ★难读"),
        ("individualized", "/ˌɪndɪˈvɪdʒuəlaɪzd/", "个性化的"),
        ("distress", "/dɪˈstres/", "痛苦/困扰"),
    ],
    "练习7 社区姑息护理中护理角色的提升": [
        ("palliative", "/ˈpæliətɪv/", "姑息的/安宁的 ★核心词"),
        ("holistic", "/həˈlɪstɪk/", "整体的"),
        ("fragmented", "/ˈfræɡmentɪd/", "碎片化的"),
        ("documented", "/ˈdɒkjumentɪd/", "有文献记录的"),
        ("uneven", "/ʌnˈiːvn/", "不均衡的"),
        ("neglects", "/nɪˈɡlekts/", "忽视"),
    ],
    "练习8 从职业倦怠到韧性": [
        ("disruption", "/dɪsˈrʌpʃn/", "中断/破坏"),
        ("morally", "/ˈmɒrəli/", "道德上地"),
        ("longitudinal", "/ˌlɒndʒɪˈtjuːdɪnl/", "纵向的 ★难读"),
        ("threshold", "/ˈθreʃhəʊld/", "阈值 ★θ+ʃ连读"),
        ("depletion", "/dɪˈpliːʃn/", "耗竭"),
        ("durability", "/ˌdjʊərəˈbɪləti/", "持久性"),
        ("escalated", "/ˈeskəleɪtɪd/", "急剧加重"),
    ],
    "练习9 护士处方权": [
        ("bibliometric", "/ˌbɪbliəˈmetrɪk/", "文献计量的 ★学术词"),
        ("delegating", "/ˈdelɪɡeɪtɪŋ/", "授权"),
        ("accessibility", "/əkˌsesəˈbɪləti/", "可及性"),
        ("prevalence", "/ˈprevələns/", "流行率/患病率 ★高频"),
        ("subsequently", "/ˈsʌbsɪkwəntli/", "随后"),
        ("prescribed", "/prɪˈskraɪbd/", "开具处方的"),
    ],
    "练习10 物联网与智慧护理": [
        ("unprecedented", "/ʌnˈpresɪdentɪd/", "前所未有的 ★难读"),
        ("integrating", "/ˈɪntɪɡreɪtɪŋ/", "整合"),
        ("substantially", "/səbˈstænʃəli/", "显著地"),
        ("surveillance", "/sɜːˈveɪləns/", "监测/监控"),
        ("demographics", "/ˌdeməˈɡræfɪks/", "人口统计"),
    ],
    "练习11 循证护理实践": [
        ("fundamental", "/ˌfʌndəˈmentl/", "基本的"),
        ("optimize", "/ˈɒptɪmaɪz/", "优化"),
        ("proficiency", "/prəˈfɪʃənsi/", "熟练度"),
        ("baccalaureate", "/ˌbækəˈlɔːriət/", "学士学位 ★难读"),
        ("methodology", "/ˌmeθəˈdɒlədʒi/", "方法论"),
        ("progression", "/prəˈɡreʃn/", "进展/提升"),
    ],
    "练习12 用药错误与患者安全": [
        ("mortality", "/mɔːˈtæləti/", "死亡率 ★高频"),
        ("morbidity", "/mɔːˈbɪdəti/", "发病率 ★与mortality对比记"),
        ("encompassing", "/ɪnˈkʌmpəsɪŋ/", "涵盖"),
        ("dispensing", "/dɪˈspensɪŋ/", "调配（药物）"),
        ("inherently", "/ɪnˈherəntli/", "固有地/本质上"),
        ("vulnerable", "/ˈvʌlnərəbl/", "脆弱的/易受伤害的"),
        ("verification", "/ˌverɪfɪˈkeɪʃn/", "核验"),
    ],
    "3.18面试 职业暴露": [
        ("occupational", "/ˌɒkjuˈpeɪʃənl/", "职业的"),
        ("exposure", "/ɪkˈspəʊʒə/", "暴露"),
        ("lancet", "/ˈlɑːnsɪt/", "采血针/柳叶刀 ★易忘"),
        ("ampoule", "/ˈæmpuːl/", "安瓿（玻璃药瓶）★易忘"),
        ("specificity", "/ˌspesɪˈfɪsəti/", "特殊性"),
        ("conducive", "/kənˈdjuːsɪv/", "有助于的/导致的"),
        ("haste", "/heɪst/", "仓促/急躁"),
        ("eliminate", "/ɪˈlɪmɪneɪt/", "消除"),
        ("pricking", "/ˈprɪkɪŋ/", "刺穿"),
    ],
}

# ══════════════════════════════════════════════════════════
#  PART 2 : 补充朗读练习
# ══════════════════════════════════════════════════════════

NEW_READINGS = [
    {
        "title": "练习13: 护士职业暴露与工作场所安全",
        "source": "Occupational hazards and safety practices among nursing professionals. Journal of Nursing Management, 2025.",
        "en": """Healthcare workers, particularly nurses, face significant occupational hazards in their daily practice. Needle stick injuries, exposure to infectious body fluids, and contact with hazardous medications constitute the primary biological and chemical risks. The World Health Organization estimates that approximately 2 million needle stick injuries occur annually among healthcare workers worldwide, with nurses accounting for the majority of cases.

Preventive strategies encompass multiple dimensions. Standard precautions — including proper hand hygiene, appropriate use of personal protective equipment (PPE), and safe handling of sharps — form the cornerstone of infection prevention. Additionally, engineering controls such as safety-engineered needles and needleless intravenous systems have demonstrated significant reductions in percutaneous injuries. However, implementation barriers persist, including insufficient training, time pressure during emergencies, and inadequate reporting mechanisms.

Post-exposure protocols require immediate wound management, risk assessment based on the source patient's infectious status, and timely prophylactic treatment when indicated. Establishing a non-punitive reporting culture is essential, as underreporting — estimated at 40-70% globally — severely undermines institutional capacity to identify patterns and implement corrective measures.""",
        "zh": """医疗工作者，尤其是护士，在日常工作中面临重大的职业危害。针刺伤、感染性体液暴露和接触有害药物构成主要的生物和化学风险。世界卫生组织估计，全球医疗工作者每年约发生200万次针刺伤，其中护士占大多数。

预防策略涵盖多个维度。标准预防——包括正确的手卫生、适当使用个人防护装备（PPE）以及安全处理锐器——构成感染预防的基石。此外，工程控制措施如安全型针头和无针静脉输液系统已显著降低了经皮损伤的发生率。然而，实施障碍依然存在，包括培训不足、紧急情况下的时间压力以及报告机制不完善。

暴露后处理方案要求立即进行伤口处理、根据源患者的感染状况进行风险评估，以及在有指征时及时进行预防性治疗。建立非惩罚性报告文化至关重要，因为漏报（全球估计达40-70%）严重削弱了机构识别规律和实施纠正措施的能力。""",
        "vocab": [
            ("hazards", "/ˈhæzədz/", "危害"),
            ("percutaneous", "/ˌpɜːkjuˈteɪniəs/", "经皮的 ★难读"),
            ("prophylactic", "/ˌprɒfɪˈlæktɪk/", "预防性的 ★难读"),
            ("cornerstone", "/ˈkɔːnəstəʊn/", "基石"),
            ("non-punitive", "/nɒn ˈpjuːnɪtɪv/", "非惩罚性的"),
            ("underreporting", "/ˌʌndərɪˈpɔːtɪŋ/", "漏报"),
            ("corrective", "/kəˈrektɪv/", "纠正的"),
            ("hygiene", "/ˈhaɪdʒiːn/", "卫生"),
        ],
    },
    {
        "title": "练习14: 老龄化社会与长期护理",
        "source": "Challenges and innovations in long-term care nursing for aging populations. Geriatric Nursing, 2025.",
        "en": """The acceleration of global population aging presents profound challenges for healthcare systems. China has entered a moderately aged society, with over 300 million citizens aged 60 and above. Among them, approximately 45 million are classified as disabled or cognitively impaired, requiring varying degrees of long-term nursing care. This demographic shift demands fundamental transformation in care delivery models.

Long-term care insurance, recently designated as China's "sixth social insurance," represents a landmark policy response. By late 2025, the program covered nearly 300 million people and had benefited over 3.3 million disabled individuals. The policy aims to bridge the critical gap between escalating care needs and available resources, addressing both the 330,000-bed shortage and the 1.61 million nursing workforce deficit.

Community-based integrated care models — combining medical treatment, rehabilitation, and daily living assistance — are emerging as the predominant paradigm. Nurses in these settings serve as coordinators, educators, and advocates, conducting comprehensive geriatric assessments, managing multiple chronic conditions simultaneously, and facilitating seamless transitions between hospital, community, and home care environments.""",
        "zh": """全球人口老龄化的加速为医疗体系带来了深远的挑战。中国已进入中度老龄化社会，60岁及以上公民超过3亿。其中约4500万人为失能或认知障碍人群，需要不同程度的长期护理照护。这一人口结构变化要求护理服务模式进行根本性转型。

长期护理保险——近期被定位为中国的"第六项社会保险"——代表着一项里程碑式的政策回应。截至2025年底，该计划覆盖近3亿人，惠及超330万失能群众。该政策旨在弥合不断增长的护理需求与可用资源之间的关键缺口，同时应对33万张床位缺口和161万护理人员缺口。

社区为基础的医养结合模式——融合医疗、康复和日常生活照料——正成为主流范式。护士在这些环境中担任协调者、教育者和倡导者的角色，开展老年综合评估、同时管理多种慢性疾病，并促进医院、社区和居家护理环境之间的无缝衔接。""",
        "vocab": [
            ("cognitively", "/ˈkɒɡnɪtɪvli/", "认知上地"),
            ("impaired", "/ɪmˈpeəd/", "受损的/障碍的"),
            ("demographic", "/ˌdeməˈɡræfɪk/", "人口统计的"),
            ("designated", "/ˈdezɪɡneɪtɪd/", "被指定为"),
            ("predominant", "/prɪˈdɒmɪnənt/", "主要的/占主导的"),
            ("paradigm", "/ˈpærədaɪm/", "范式 ★难读"),
            ("geriatric", "/ˌdʒeriˈætrɪk/", "老年医学的 ★常考"),
            ("rehabilitation", "/ˌriːəˌbɪlɪˈteɪʃn/", "康复"),
            ("seamless", "/ˈsiːmləs/", "无缝的"),
        ],
    },
]

# ══════════════════════════════════════════════════════════
#  PART 3 : 补充模拟面试
# ══════════════════════════════════════════════════════════

NEW_MOCK_INTERVIEWS = [
    {
        "title": "模拟面试9：长期护理保险与医养结合",
        "reading_en": """The integration of medical care and elderly support represents a strategic response to population aging in China. The State Council's guidance on promoting medical-nursing integration emphasizes establishing a coordinated system encompassing home-based, community-based, and institutional care. Healthcare institutions are encouraged to embed nursing services within elderly care facilities, while community health centers extend their reach through mobile nursing teams and telehealth consultations.

Despite significant policy momentum, implementation challenges remain substantial. Fragmented administrative oversight — with healthcare under the National Health Commission and elderly care under the Ministry of Civil Affairs — creates coordination difficulties. Additionally, the current nursing workforce lacks adequate geriatric training, and reimbursement mechanisms fail to incentivize integrated service delivery adequately.""",
        "reading_zh": """医疗与养老服务的融合是中国应对人口老龄化的战略举措。国务院关于促进医养结合的指导意见强调建立涵盖居家、社区和机构的协调体系。鼓励医疗机构在养老设施中嵌入护理服务，同时社区卫生服务中心通过流动护理团队和远程健康咨询扩展服务范围。

尽管政策推动力度显著，实施挑战依然严峻。行政管理碎片化——医疗由国家卫健委管辖，养老由民政部管辖——造成了协调困难。此外，当前护理队伍缺乏充分的老年护理培训，报销机制也未能充分激励整合式服务的提供。""",
        "reading_vocab": [
            ("embed", "/ɪmˈbed/", "嵌入"),
            ("reimbursement", "/ˌriːɪmˈbɜːsmənt/", "报销 ★难读"),
            ("incentivize", "/ɪnˈsentɪvaɪz/", "激励"),
            ("fragmented", "/ˈfræɡmentɪd/", "碎片化的"),
            ("oversight", "/ˈəʊvəsaɪt/", "监管/监督"),
            ("momentum", "/məˈmentəm/", "势头/动力"),
        ],
        "question": "请谈谈你对长期护理保险制度和医养结合模式的看法。护士在其中扮演什么角色？",
        "answer": """各位老师好，非常感谢这个问题。我将从"长期护理保险的意义"和"医养结合中护士的角色"两个方面展开。

第一部分：长期护理保险制度的重要意义

长期护理保险被称为继养老、医疗、工伤、失业、生育之后的"社保第六险"，是我国应对深度老龄化社会的关键制度安排。截至2025年底，长护险已覆盖近3亿人，惠及超330万失能群众。

在临床中，我深切感受到长护险的迫切需要。泌尿外科的许多老年患者——前列腺增生术后需要长期留置导尿管的、膀胱癌术后行泌尿造口的——出院后面临巨大的居家护理负担。过去，很多家庭因经济原因不得不放弃专业的延续性护理，导致并发症频发。长护险通过制度化保障，让这些患者能够获得持续的专业照护，真正实现了"有人照料、有钱护理"。

当然，长护险也面临挑战：一是护理人员缺口巨大——目前约161万人；二是服务标准尚不统一；三是筹资可持续性需要进一步验证。

第二部分：医养结合中护士的角色

医养结合的核心理念是打破医疗与养老的壁垒，构建"居家-社区-机构"三位一体的服务体系。在这个体系中，护士扮演着不可替代的枢纽角色：

第一，综合评估者。通过老年综合评估（CGA），全面评估老年人的身体功能、认知状态、心理状况、社会支持和居家环境，为制定个性化护理方案提供依据。

第二，多病共存管理者。老年患者往往同时患有高血压、糖尿病、前列腺增生等多种慢性病，需要护士具备跨学科的综合管理能力，协调用药方案、监测药物相互作用。

第三，连续性照护协调者。在患者从医院出院到社区、再到居家的过程中，护士负责制定出院计划、进行交接沟通、开展远程随访，确保护理的连续性不被打断。

第四，健康教育者和家属培训者。教会家属基本的护理技能——如造口护理、导管维护、翻身拍背等——是减轻照护负担、预防并发症的关键。

在我的临床实践中，我曾护理过一位78岁的肾造瘘术后患者，出院后需要长期带管。他的老伴年迈体弱，独生子女在外地工作。通过医养结合的社区护理服务，我们安排了定期上门换管和远程随访指导，同时帮助其申请了长护险补贴。三个月后复查时，老人的造瘘管维护良好，无一次感染。这个案例让我深刻认识到：好的制度设计加上专业的护理执行，才能真正解决老龄化社会的照护难题。

这也是我选择攻读研究生的重要原因之一——我希望系统学习老年护理和慢病管理的前沿理论，为即将到来的深度老龄化社会做好准备。

我的回答完毕，谢谢各位老师！""",
    },
    {
        "title": "模拟面试10：安宁疗护与工作场所安全",
        "reading_en": """Hospice care, also known as end-of-life care, focuses on providing comprehensive comfort and support to patients with terminal illnesses and their families. Unlike curative treatment, hospice care prioritizes quality of life, symptom management, and psychological and spiritual support during the final stages of life. China's National Health Commission issued the updated Practice Guidelines for Hospice Care in 2025, marking a significant step toward standardizing end-of-life care services nationwide.

Despite growing policy support, hospice care development in China faces considerable obstacles. Cultural attitudes toward death — deeply influenced by Confucian and Buddhist traditions — often result in reluctance to discuss end-of-life planning. Furthermore, a severe shortage of trained hospice care professionals, limited public awareness, and insufficient insurance coverage continue to constrain service accessibility.""",
        "reading_zh": """安宁疗护，也称临终关怀，专注于为终末期疾病患者及其家庭提供全面的舒适照护与支持。与治愈性治疗不同，安宁疗护优先关注生命末期的生活质量、症状管理以及心理和灵性支持。2025年，国家卫健委发布了更新版的《安宁疗护实践指南》，标志着在全国范围内规范临终关怀服务迈出了重要一步。

尽管政策支持力度不断增强，安宁疗护在中国的发展仍面临相当大的障碍。受儒家和佛教传统深刻影响的死亡观念，常常导致人们不愿讨论临终规划。此外，受过培训的安宁疗护专业人员严重短缺、公众认知有限以及保险覆盖不足，持续制约着服务的可及性。""",
        "reading_vocab": [
            ("hospice", "/ˈhɒspɪs/", "安宁疗护/临终关怀 ★核心词"),
            ("curative", "/ˈkjʊərətɪv/", "治愈性的"),
            ("terminal", "/ˈtɜːmɪnl/", "终末期的"),
            ("spiritual", "/ˈspɪrɪtʃuəl/", "灵性的/精神的"),
            ("Confucian", "/kənˈfjuːʃən/", "儒家的"),
            ("reluctance", "/rɪˈlʌktəns/", "不情愿"),
            ("constrain", "/kənˈstreɪn/", "制约/限制"),
        ],
        "question": "请谈谈你对安宁疗护在中国推广的看法。推广面临哪些障碍？护士应该发挥什么作用？",
        "answer": """各位老师好，我的回答将从"安宁疗护的认知"、"推广障碍"和"护士的角色"三个方面展开。

第一部分：对安宁疗护的理解

安宁疗护的核心理念是：当治愈不再可能时，让患者有尊严地、舒适地走完人生最后一程。WHO将其定义为"通过早期识别、准确评估和治疗疼痛及其他身体、心理、社会和灵性问题，预防和缓解痛苦，从而改善面临危及生命疾病的患者及其家庭的生活质量"。

2025年国家卫健委印发的《安宁疗护实践指南（2025年版）》，标志着我国安宁疗护从"理念倡导"走向"规范实施"的重要转折。指南明确了以多学科团队为核心、以患者和家属为中心的服务模式。

我想特别强调：安宁疗护不等于"放弃治疗"。它是一种积极的、有技术含量的护理——疼痛管理、症状控制、心理疏导、灵性关怀，每一项都需要专业知识和技能。

第二部分：推广面临的障碍

第一，文化障碍。中国传统文化中"死亡禁忌"根深蒂固——很多家属认为谈论死亡"不吉利"，甚至觉得选择安宁疗护就是"不孝"。在泌尿外科，我遇到过晚期膀胱癌患者的家属，明知治愈无望，仍坚持要求"再做一次化疗"，因为"不能让别人觉得我们放弃了老人"。

第二，专业人才短缺。安宁疗护需要系统的疼痛管理、心理咨询、沟通技巧培训，但目前我国护理本科和硕士课程中，安宁疗护相关内容非常有限。

第三，制度保障不足。安宁疗护的医保覆盖范围有限，很多心理疏导、灵性照护等非技术性服务无法报销，导致患者经济负担重。

第四，公众认知不够。很多人不知道安宁疗护的存在，更不了解它能提供什么。

第三部分：护士在安宁疗护中的角色

第一，症状管理专家。疼痛是终末期患者最常见也最恐惧的症状。护士需掌握NRS疼痛评估、WHO三阶梯镇痛原则，确保"让患者无痛或尽可能少痛"。

第二，心理支持者。通过叙事护理的方法——倾听患者的生命故事、帮助其回顾人生意义、表达未尽的心愿——给予深层次的心理关怀。

第三，家属教育者和哀伤辅导者。帮助家属理解病情进展、做好心理准备，在患者离世后提供哀伤辅导，预防病理性哀伤。

第四，生死教育的推动者。护士可以通过社区讲座、公众号科普等方式，传播"生死两安"的理念，逐步改变社会对死亡的认知。

在泌尿外科，我曾照护一位晚期肾癌的老教授。最初他非常抗拒，觉得"我还没到那一步"。后来随着病情进展，剧烈的骨转移疼痛让他彻夜难眠。在安宁疗护团队介入后，通过规范的镇痛方案和持续的心理陪伴，他最后的日子是安详的。他的女儿后来对我说："谢谢你们让爸爸走得没有那么痛苦。"这句话让我至今记忆深刻——这就是安宁疗护的意义。

这也是我希望在研究生阶段深入学习肿瘤护理的重要原因。我相信，随着制度完善和观念转变，安宁疗护将在中国迎来更广阔的发展空间。

我的回答完毕，谢谢各位老师！""",
    },
]

# ══════════════════════════════════════════════════════════
#  PART 4 : 速记模板更新
# ══════════════════════════════════════════════════════════

NEW_HOTSPOTS = [
    ("长期护理保险(社保第六险)", "覆盖近3亿人,十五五从试点转全面建制,161万护理人员缺口", "老龄化、社区护理"),
    ("安宁疗护实践指南(2025年版)", "国家卫健委印发,规范终末期照护,多学科团队+患者家属中心", "姑息/安宁护理"),
    ("医养结合促进行动(2025)", "居家-社区-机构三位一体,医-养-康-护融合", "老龄化、社区护理"),
    ("护理工作场所暴力", "70.8%护士遭遇过,离职意向率42.8%,WHO发布应对指南", "职业安全、倦怠"),
    ("AI+Nursing国自然重点方向(2026)", "人机协同、数字孪生、智能照护、AI护理决策支持", "AI护理"),
    ("全球护士短缺新数据", "WHO 2025报告:护士增长但不平等加剧,预计2030年缺口410万", "人力、政策"),
]

NEW_DATA = [
    ("长护险覆盖近3亿人，惠及超330万失能群众", "国家医保局2025", "长护险、老龄化"),
    ("中国失能失智老人超4500万", "国家卫健委", "长护险、老龄化"),
    ("护理人员缺口161万，床位缺口33万张", "国家卫健委", "长护险、人力"),
    ("中国护士离职意向率42.8%", "六省调查2025", "倦怠、人力"),
    ("70.8%护士遭遇过工作场所暴力", "中国护理调查", "职业安全"),
    ("互联网+护理行业规模突破500亿", "2025行业报告", "远程护理"),
    ("全国注册护士达563万，医护比1:1.17", "国家卫健委2025", "人力、政策"),
    ("安宁疗护实践指南（2025年版）印发", "国家卫健委", "安宁疗护"),
]

NEW_FOLLOWUPS = [
    ("你对长期护理保险制度怎么看？",
     "「长期护理保险是我国应对老龄化的关键制度创新，被称为'社保第六险'。截至2025年，已覆盖近3亿人。作为泌尿外科护士，我深感其必要性——很多术后需长期带管的老年患者，因经济原因放弃了专业的延续性护理。长护险的推广，将使'专业居家照护'从奢侈品变为基本保障。当然，推广中需解决护理人员缺口（约161万）、服务标准统一、筹资可持续性等问题。这也是我读研的动力之一——希望为长护险体系下的专业护理服务贡献力量。」"),

    ("你如何看待护理工作场所暴力？",
     "「护理工作场所暴力是全球性问题。数据显示，中国约70.8%护士遭遇过不同形式的暴力，包括言语侮辱、肢体攻击等，这是导致离职意向率高达42.8%的重要因素。应对需多管齐下：一是完善法律保护，明确医疗场所暴力的法律后果；二是加强安保措施，建立一键报警系统；三是改善沟通，很多暴力源于信息不对称和等候过长；四是建立心理支持，为受害护士提供咨询和法律援助。在泌尿外科，我们也遇到过情绪激动的家属——关键是保持专业冷静，用SBAR沟通法清晰说明情况。」"),

    ("谈谈你对安宁疗护在中国推广的看法。",
     "「安宁疗护不是'放弃治疗'，而是'以舒适为目标的全人照护'。2025年国家卫健委印发《安宁疗护实践指南》标志着从'理念倡导'走向'规范实施'。推广面临三大障碍：文化上的'死亡禁忌'、专业人才短缺、医保覆盖不足。护士应成为症状管理专家、心理支持者、家属哀伤辅导者和生死教育推动者。在泌尿外科，晚期肿瘤患者的安宁疗护需求很大——不仅控制疼痛，更要帮助患者和家属找到生命的意义。治愈有时，帮助常常，安慰总是。」"),

    ("你对医养结合模式中护士角色的理解？",
     "「医养结合的核心是打破医疗与养老的壁垒。护士在其中扮演四重角色：综合评估者——通过老年综合评估（CGA）全面了解老人需求；多病共存管理者——协调多种慢病的用药和护理方案；连续性照护协调者——确保医院-社区-居家的无缝衔接；健康教育者——培训家属基本护理技能。面临的挑战包括行政管理碎片化（卫健委vs民政部）、护士老年护理培训不足、报销机制不完善等。这正是需要培养更多高学历护理人才的原因。」"),
]

# ══════════════════════════════════════════════════════════
#  PART 5 : 一页纸速记卡
# ══════════════════════════════════════════════════════════

QUICK_REF = {
    "framework": "概-例-策-升：概念+数据(25%) → 泌尿外科故事(35%) → 3条对策(25%) → 读研动机(15%)",
    "topics": [
        ("AI护理", "工具不替代，共情无法编程"),
        ("人文关怀", "护理灵魂，技术时代更要主动填充温度"),
        ("循证护理", "从「凭经验」到「凭证据」，5A循环"),
        ("肿瘤/姑息", "不是放弃，是以舒适为目标的全人照护"),
        ("处方权", "小范围·高标准·严监管（法-资-目-监-保）"),
        ("职业倦怠", "结构性问题，非个人软弱（Maslach三维度）"),
        ("远程护理", "最后一公里解决方案，数字鸿沟摆渡人"),
        ("用药安全", "非惩罚报告+系统改进，三查八对不可逾越"),
        ("长护险+医养", "第六险，3亿覆盖，161万缺口，四重角色"),
        ("安宁疗护", "让生命有尊严地谢幕，2025新指南"),
    ],
    "theories": [
        ("Watson", "10大关怀因子，人文价值观→关怀行为"),
        ("叙事护理", "关注→再现→归属（Charon 2001）"),
        ("循证5A", "Ask→Acquire→Appraise→Apply→Assess"),
        ("Maslach", "情感耗竭→去人格化→成就感↓"),
        ("时效性激励", "按时间节点差异化干预（自己论文）"),
        ("健康中国2030", "以治疗为中心→以健康为中心"),
    ],
    "data": [
        "全球护士缺口2030: 410万 | 中国护士563万",
        "倦怠率62%，<25岁75% | 离职意向42.8%",
        "用药错误年损失420亿$ | 1/3在配制给药阶段",
        "长护险覆盖3亿，缺口161万 | 失能老人4500万",
        "早期姑息↑生活质量11分 | 远程监测率26%",
        "工作场所暴力70.8% | 需姑息护理4000万人/年",
        "互联网+护理规模500亿 | 65+人口2050→16亿",
    ],
    "quotes": [
        "算法算概率，不能替代床畔安抚的眼神",
        "技术监测尿量，无法丈量内心哀伤",
        "治愈有时，帮助常常，安慰总是",
        "从「我觉得」到「研究证明」",
        "让有能力的人做有意义的事",
        "照顾好自己，才能照顾好别人",
        "让专业护理跨越医院围墙",
        "每条安全制度背后都有沉痛教训",
    ],
}


# ══════════════════════════════════════════════════════════
#  BUILD DOCUMENTS
# ══════════════════════════════════════════════════════════

def build_main_doc():
    doc = Document()
    init_doc_styles(doc)

    # ── 标题页 ───────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('护理复试材料补充包', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, cn_font=CN_HEADING)
    sub = doc.add_paragraph('南京大学护理学硕士复试 · 汪妍专用')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_run_font(sub.runs[0])

    date_p = doc.add_paragraph('2026年3月 · 补充版')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(12)
    date_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    set_run_font(date_p.runs[0])

    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc.paragraph_format.space_before = Pt(36)
    toc_items = [
        "第一部分  英文朗读难词标注速查表",
        "第二部分  补充朗读练习（练习13-14）",
        "第三部分  补充模拟面试（第9-10套）",
        "第四部分  速记模板更新（热点·数据·追问）",
        "第五部分  一页纸速记卡",
    ]
    for item in toc_items:
        run = toc.add_run(item + "\n")
        run.font.size = Pt(11)
        set_run_font(run)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 第一部分：难词标注
    # ═══════════════════════════════════════════════
    add_heading(doc, '第一部分  英文朗读难词标注速查表', level=1)
    add_para(doc, '说明：按练习编号排列，标注每篇中容易读错/卡壳的专业词汇。★标记为高频或特别难读的词。打印后建议用荧光笔标记自己不熟悉的词，每天朗读前先过一遍。', size=9)

    for section_name, words in VOCAB_DATA.items():
        add_heading(doc, section_name, level=3)
        add_table(doc,
                  ["英文", "音标", "中文释义"],
                  [(w[0], w[1], w[2]) for w in words],
                  col_widths=[4.5, 5.5, 6])
        doc.add_paragraph()  # spacing

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 第二部分：补充朗读练习
    # ═══════════════════════════════════════════════
    add_heading(doc, '第二部分  补充朗读练习', level=1)
    add_para(doc, '说明：补充2篇练习，覆盖3.18面试涉及的"职业暴露"主题以及2026年热点"老龄化与长期护理"主题。格式与原12篇一致，每篇约200-300词，限时5分钟朗读翻译。', size=9)

    for reading in NEW_READINGS:
        add_heading(doc, reading["title"], level=2)
        add_para(doc, f'来源: {reading["source"]}', size=9)

        add_heading(doc, '【英文原文】', level=4)
        for para in reading["en"].strip().split("\n\n"):
            add_para(doc, para.strip(), size=10.5)

        add_heading(doc, '【中文翻译】', level=4)
        for para in reading["zh"].strip().split("\n\n"):
            add_para(doc, para.strip(), size=10.5)

        add_heading(doc, '【本篇难词标注】', level=4)
        add_table(doc,
                  ["英文", "音标", "中文释义"],
                  [(w[0], w[1], w[2]) for w in reading["vocab"]],
                  col_widths=[4.5, 5.5, 6])
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 第三部分：补充模拟面试
    # ═══════════════════════════════════════════════
    add_heading(doc, '第三部分  补充模拟面试', level=1)
    add_para(doc, '说明：补充2套完整模拟面试，覆盖"长护险+医养结合"和"安宁疗护"两个2026年最新热点。格式与原8套一致：朗读翻译段落 + 专业问答。', size=9)

    for mock in NEW_MOCK_INTERVIEWS:
        add_heading(doc, mock["title"], level=2)

        # 朗读段落
        add_para(doc, '请朗读并翻译以下文章，限时5分钟。', bold=True, size=10.5)
        for para in mock["reading_en"].strip().split("\n\n"):
            add_para(doc, para.strip(), size=10.5)

        # 翻译
        for para in mock["reading_zh"].strip().split("\n\n"):
            add_para(doc, para.strip(), size=10.5)

        # 难词
        add_heading(doc, '【朗读段落难词】', level=4)
        add_table(doc,
                  ["英文", "音标", "中文释义"],
                  [(w[0], w[1], w[2]) for w in mock["reading_vocab"]],
                  col_widths=[4.5, 5.5, 6])
        doc.add_paragraph()

        # 问题
        add_para(doc, mock["question"], bold=True, size=11)

        # 回答
        for para in mock["answer"].strip().split("\n\n"):
            add_para(doc, para.strip(), size=10.5)

        doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 第四部分：速记模板更新
    # ═══════════════════════════════════════════════
    add_heading(doc, '第四部分  速记模板更新', level=1)
    add_para(doc, '说明：以下内容为原《答题套路速记模板》的增量更新，打印后夹入原文档对应位置即可。', size=9)

    # 4.1 新增热点
    add_heading(doc, '4.1 新增热点时事（补充到"第四节 热点时事"表格中）', level=2)
    add_table(doc,
              ["热点", "要点", "可用话题"],
              NEW_HOTSPOTS,
              col_widths=[4.5, 7.5, 4])
    doc.add_paragraph()

    # 4.2 新增数据
    add_heading(doc, '4.2 新增数据弹药（补充到"第七节 数据弹药库"中）', level=2)
    add_table(doc,
              ["数据", "来源", "适用话题"],
              NEW_DATA,
              col_widths=[7, 4.5, 4.5])
    doc.add_paragraph()

    # 4.3 新增话题速记
    add_heading(doc, '4.3 新增话题核心论点速记', level=2)

    add_heading(doc, '话题9：长期护理保险与医养结合', level=3)
    add_para(doc, '一句话定位：长护险是应对老龄化的"第六道防线"，医养结合是打通医疗与养老的"最后一公里"。', bold=True)
    add_para(doc, """必说论点：
1. 为什么急需：中国失能失智老人超4500万，护理人员缺口161万，床位缺口33万张。很多家庭"有心照料，无力承担"。
2. 长护险价值：覆盖近3亿人，让专业照护从"自费奢侈品"变为"制度化保障"。泌尿外科大量术后长期带管患者直接受益。
3. 医养结合中护士四重角色：综合评估者（CGA）+ 多病共存管理者 + 连续性照护协调者 + 健康教育者。""")
    add_para(doc, """泌尿外科故事：
> 一位78岁肾造瘘术后患者，老伴年迈、子女在外地。通过医养结合社区服务安排定期上门换管+远程随访，申请长护险补贴后三个月无一次感染。好制度+好护理=真正的老有所护。""")
    add_para(doc, '收尾金句：「让每一位老人都能在熟悉的环境中获得专业的照护——这是长护险的初心，也是护理的使命。」', bold=True)

    add_heading(doc, '话题10：安宁疗护', level=3)
    add_para(doc, '一句话定位：安宁疗护不是"放弃"，而是让生命有尊严地谢幕。', bold=True)
    add_para(doc, """必说论点：
1. 核心理念：WHO定义——预防和缓解痛苦，改善终末期患者及家庭的生活质量。2025年《安宁疗护实践指南》标志"理念→规范"转折。
2. 四大障碍：文化禁忌（谈死不吉利）、人才短缺、医保覆盖不足、公众认知有限。
3. 护士四重角色：症状管理专家（NRS+三阶梯镇痛）+ 心理支持者（叙事护理）+ 哀伤辅导者 + 生死教育推动者。""")
    add_para(doc, """泌尿外科故事：
> 一位晚期肾癌老教授，从抗拒姑息到接受安宁疗护。规范镇痛+心理陪伴让他最后的日子安详平静。女儿说："谢谢你们让爸爸走得没那么痛苦。"——这就是安宁疗护的意义。""")
    add_para(doc, '收尾金句：「生命的终点不是医学的失败，而是人文关怀最深刻的体现。」', bold=True)
    doc.add_paragraph()

    # 4.4 新增追问
    add_heading(doc, '4.4 新增追问完整应答（补充到"第八节 常见追问"中）', level=2)
    for q, a in NEW_FOLLOWUPS:
        add_heading(doc, f'"{q}"', level=3)
        add_para(doc, a, size=10.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 第五部分：一页纸速记卡
    # ═══════════════════════════════════════════════
    add_heading(doc, '第五部分  一页纸速记卡', level=1)
    add_para(doc, '★ 考前30分钟快速过一遍 ★', bold=True, size=12)

    # 框架
    add_heading(doc, '万能框架', level=3)
    add_para(doc, QUICK_REF["framework"], bold=True, size=10)

    # 10大话题
    add_heading(doc, '10大话题·一句话定位', level=3)
    add_table(doc,
              ["#", "话题", "一句话定位"],
              [(str(i+1), t[0], t[1]) for i, t in enumerate(QUICK_REF["topics"])],
              col_widths=[1, 3, 12])
    doc.add_paragraph()

    # 理论速记
    add_heading(doc, '6大理论速记', level=3)
    add_table(doc,
              ["理论", "核心要点"],
              QUICK_REF["theories"],
              col_widths=[3.5, 12.5])
    doc.add_paragraph()

    # 数据
    add_heading(doc, '核心数据（7组）', level=3)
    for d in QUICK_REF["data"]:
        add_para(doc, f'• {d}', size=10, space_after=2)

    # 金句
    add_heading(doc, '金句库（8句）', level=3)
    for i, q in enumerate(QUICK_REF["quotes"]):
        add_para(doc, f'{i+1}. 「{q}」', size=10, space_after=2)

    # 应急
    add_heading(doc, '应急锦囊', level=3)
    add_para(doc, '不懂的题 → "虽然我研究有限，但结合五年泌尿外科临床体会…" → 关联熟悉领域 → "这是我读研想深入的方向"', size=10)
    add_para(doc, '紧张时 → 深呼吸3秒 → "各位老师好，非常荣幸回答这个问题" → 宣布结构（两/三方面）→ 按步骤展开', size=10)

    # save
    path = os.path.join(OUT_DIR, '复试材料补充包-完整版.docx')
    doc.save(path)
    print(f'✅ 已生成: {path}')
    return path


if __name__ == '__main__':
    build_main_doc()
    print('\n🎉 补充材料生成完毕！')
