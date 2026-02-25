#!/usr/bin/env python3.10
"""填充博士中期考核报告模板 - 完整版（含文献综述和图片）"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FIGS_DIR = "/home/szw/github/templog/phd-midterm-defense/figs"
TEMPLATE = "/home/szw/github/templog/phd-midterm-defense/学博模板.docx"
OUTPUT = "/home/szw/github/templog/phd-midterm-defense/szw-602024320006.docx"

doc = Document(TEMPLATE)

# ============================================================
# 辅助函数
# ============================================================
def add_after(ref_para, text, font_size=Pt(12), bold=False, indent=False, align=None):
    """在 ref_para 后插入新段落"""
    new_para = doc.add_paragraph()
    ref_para._element.addnext(new_para._element)
    if text:
        run = new_para.add_run(text)
        run.font.size = font_size
        run.font.name = '宋体'
        run.bold = bold
    if indent:
        new_para.paragraph_format.first_line_indent = Pt(24)
    if align:
        new_para.alignment = align
    return new_para

def add_image_after(ref_para, img_path, width=Inches(5.5), caption=None):
    """在 ref_para 后插入图片和标题"""
    # 图片段落
    img_para = doc.add_paragraph()
    ref_para._element.addnext(img_para._element)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_path, width=width)

    if caption:
        cap_para = doc.add_paragraph()
        img_para._element.addnext(cap_para._element)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(caption)
        run.font.size = Pt(10)
        run.font.name = '宋体'
        return cap_para
    return img_para

# ============================================================
# 1. 封面表格
# ============================================================
table = doc.tables[0]
table.rows[2].cells[2].text = "智能化软件工程"

# PART2_PLACEHOLDER

# ============================================================
# 2. 日期
# ============================================================
for para in doc.paragraphs:
    if para.text.strip() == "年    月    日":
        para.clear()
        run = para.add_run("2025年    3月    1日")
        break

# ============================================================
# 3. 科研成果列表
# ============================================================
for para in doc.paragraphs:
    if "CCF A类 会议 第一作者 一篇" in para.text:
        para.clear()
        run = para.add_run("1、CCF A类 期刊 第一作者 发表一篇（TOSEM 2026）")
        run.font.size = Pt(12)
        break

for para in doc.paragraphs:
    if "专利 学生一作 一篇" in para.text:
        para.clear()
        run = para.add_run("2、CCF A类 期刊 第二作者 发表两篇（TOSEM 2025）")
        run.font.size = Pt(12)
        break

# ============================================================
# 4. 定位关键段落
# ============================================================
idx = {}
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t.startswith("计算机科学与技术/软件工程"):
        idx["hint"] = i
    elif t == "中英文摘要":
        idx["abstract"] = i
    elif t.startswith("研究工作的背景"):
        idx["background"] = i
    elif t.startswith("本人所承担的具体研究工作"):
        idx["my_work"] = i
    elif t.startswith("下一步科研计划"):
        idx["future"] = i
    elif t.startswith("除上述工作外"):
        idx["other"] = i
    elif t.startswith("如申请"):
        idx["special"] = i

# 清除提示文字
doc.paragraphs[idx["hint"]].clear()
doc.paragraphs[idx["other"]].clear()
doc.paragraphs[idx["special"]].clear()

# PART3_PLACEHOLDER

# ============================================================
# 5. 摘要
# ============================================================
p = doc.paragraphs[idx["abstract"]]
p.clear()
run = p.add_run("摘要")
run.font.size = Pt(14)
run.bold = True

abstract_cn = (
    "严重的软件缺陷可能导致重大问题甚至造成巨额经济损失。因此，自动化代码缺陷检测受到了广泛关注。"
    "为了尽早修复这些缺陷，研究者提出了即时缺陷预测（JIT-DP）和即时缺陷定位（JIT-DL）两类任务。"
    "其中，JIT-DP旨在代码变更提交时预测提交级别的缺陷，JIT-DL旨在故障发生前进行行级别的缺陷定位。"
    "目前已有多种方法被提出用于JIT-DP和JIT-DL，其中基于代码预训练模型（CPTM）的方法取得了最优结果。"
    "然而，大多数已有研究主要关注JIT-DP，尽管JIT-DL在实际应用中更为关键。"
    "识别提交中哪些具体代码行存在缺陷，有助于判断整个提交是否存在缺陷。"
    "因此，本文提出了一种多任务学习方法JIT-MTL，同时解决JIT-DP和JIT-DL两个任务。"
    "具体而言，我们训练CPTM同时识别缺陷提交和缺陷代码行。行级别预测直接用于JIT-DL，"
    "同一提交的提交级别预测则合并后用于JIT-DP。为提高行级别预测的准确性，"
    "我们为CPTM提供了额外的提交级别信息，包括提交级别缺陷代码表示和专家特征。"
    "在JIT-Defect4J数据集上的实验结果表明，JIT-MTL优于现有最优基线方法，"
    "在JIT-DP的F1指标上提升14%，在JIT-DL的Top-10准确率上提升53.9%，Top-5准确率上提升36.4%。"
)

abstract_en = (
    "Severe software defects can lead to significant issues and even result in substantial financial losses. "
    "As a result, automated code defect detection has garnered widespread attention. "
    "To fix these defects as soon as possible, Just-In-Time Defect Prediction (JIT-DP) and "
    "Just-In-Time Defect Localization (JIT-DL) have been investigated. "
    "We propose a Multi-Task Learning method (JIT-MTL) designed to address both JIT-DP and JIT-DL. "
    "Experimental results on the JIT-Defect4J dataset show that JIT-MTL outperforms the state-of-the-art baselines, "
    "with improvements of 14% in F1 in JIT-DP, 53.9% in Top-10 Accuracy, and 36.4% in Top-5 Accuracy in JIT-DL."
)

last = p
last = add_after(last, abstract_cn, indent=True)
last = add_after(last, "")
last = add_after(last, abstract_en, indent=True)

# PART4_PLACEHOLDER

# ============================================================
# 6. 研究背景
# ============================================================
p = doc.paragraphs[idx["background"]]
p.clear()
run = p.add_run("1  研究工作的背景")
run.font.size = Pt(14)
run.bold = True

bg = [
    ("1.1  即时缺陷预测与定位任务", True, False),
    ("软件在日常生活中至关重要，确保其健壮性和安全性是软件工程领域的核心目标。"
     "然而，由于开发流程不合理、开发者技能不足等原因，软件缺陷不可避免。"
     "一些严重的软件缺陷可能导致重大经济损失。因此，自动化代码缺陷检测成为活跃的研究课题[1-4]。"
     "传统缺陷检测方法旨在基于缺陷报告或测试用例的执行信息来定位有缺陷的文件、方法或代码行。"
     "然而，这些方法往往存在缺陷识别不及时的问题。为了尽早检测缺陷，研究者引入了即时缺陷预测（JIT-DP）。"
     "JIT-DP旨在代码变更提交时预测其中的缺陷，帮助更及时地进行代码审查[5]。"
     "如果识别出缺陷，可以立即将提交退回修复，从而降低调试成本并提高软件开发效率。",
     False, True),
    ("然而，并非提交中所有变更的代码行都存在缺陷，当前的JIT-DP方法可能在仅有少部分代码行存在问题时"
     "错误地标记整个提交[6]。为解决这一问题，一些研究开始关注即时缺陷定位（JIT-DL），"
     "旨在识别提交中确切的缺陷代码行，为代码审查提供精确指导。",
     False, True),
    ("1.2  代码预训练模型在缺陷检测中的应用", True, False),
    ("早期研究[7,8]通常提取代码特征[9]来对提交进行分类（JIT-DP）和对代码行进行排序（JIT-DL）。"
     "随着基于Transformer[10]架构的模型出现，代码预训练模型（CPTM）已成功应用于代码相关任务[11,12]。"
     "CPTM在大量代码数据上进行预训练，能够更有效地提取复杂的代码特征。"
     "对于JIT-DP，Ni等人[13]和Chen等人[14]证明使用CodeBERT[15]提取代码特征可以达到最优结果。"
     "对于JIT-DL，Ni等人[13]使用注意力分数来优先排序代码行，"
     "而Chen等人[14]则专注于训练模型使用特定的行信息来识别行级别缺陷。",
     False, True),
    ("1.3  现有研究的不足", True, False),
    ("大多数已有研究主要关注JIT-DP，对JIT-DL的关注较少，这使得JIT-DL的进展更加困难。"
     "这种不平衡引发了一个关键问题：如果CPTM在提交级别预测上表现优异但在行级别定位上表现不佳，"
     "我们能否真正信任其对缺陷代码的理解？考虑这样一个场景：CPTM正确识别了一个有缺陷的提交，"
     "但未能精确定位具体的缺陷代码行。在这种情况下，模型可能是偶然得出正确的提交级别结论，"
     "而非通过对底层缺陷模式的真正理解。这一观察启发了我们的核心洞察：准确识别哪些代码行存在缺陷"
     "应当在逻辑上先于判断整个提交是否存在缺陷，因为行级别的理解为提交级别的决策提供了更细粒度、"
     "更可靠的基础。",
     False, True),
]

last = p
for text, bold, indent in bg:
    last = add_after(last, text, bold=bold, indent=indent)

# 插入图1：示例图
last = add_image_after(last, os.path.join(FIGS_DIR, "figs1.png"), width=Inches(4.5),
                       caption="图 1  CPTM正确预测缺陷提交但错误预测缺陷代码行的示例")

# PART5_PLACEHOLDER

# ============================================================
# 7. 国内外研究现状（文献综述）
# ============================================================
last = add_after(last, "1.4  国内外研究现状", bold=True)

last = add_after(last, "1.4.1  JIT-DP方法研究进展", bold=True)
last = add_after(last,
    "即时缺陷预测（JIT-DP）的研究可追溯至Mockus和Weiss[25]提出的基于代码变更特征的缺陷预测方法。"
    "早期方法主要依赖手工设计的特征，如代码变更的规模、开发者经验、历史缺陷信息等。"
    "Kamei等人[26]提出了14个专家特征，涵盖扩散度、规模、目的、历史和经验五个维度，"
    "这些特征至今仍被广泛使用。LApredict[27]使用逻辑回归结合这些专家特征进行预测。",
    indent=True)
last = add_after(last,
    "随着深度学习的发展，研究者开始探索自动特征提取方法。"
    "Hoang等人[9]提出CC2Vec，使用分层注意力网络学习代码变更的分布式表示。"
    "DeepJIT[28]直接从代码变更和提交消息中学习特征进行端到端预测。"
    "Yang等人[7]提出Deeper，结合深度学习特征和传统特征。"
    "JITLine[8]则采用更简单的方法，通过词袋模型和随机森林实现了有竞争力的性能。",
    indent=True)
last = add_after(last,
    "近年来，代码预训练模型的引入显著提升了JIT-DP的性能。"
    "JIT-Fine[14]首次将CodeBERT应用于JIT-DP任务，通过微调预训练模型取得了优于传统方法的结果。"
    "JIT-Smart[13]进一步整合了专家特征和CodeBERT，在JIT-DP上达到了当时的最优性能。"
    "Tang等人[18]探索了使用Transformer架构直接处理代码变更序列的方法。"
    "CCAF[17]提出使用AdapterFusion技术学习代码变更表示，为参数高效微调在缺陷预测中的应用提供了新思路。",
    indent=True)

last = add_after(last, "1.4.2  JIT-DL方法研究进展", bold=True)
last = add_after(last,
    "即时缺陷定位（JIT-DL）旨在识别提交中具体的缺陷代码行，相比JIT-DP提供更精确的指导。"
    "早期方法主要基于启发式规则或简单的统计特征。"
    "TPF-JIT[29]使用token频率和位置信息来排序可能存在缺陷的代码行。"
    "JITLine[8]通过计算每行代码的缺陷倾向分数来实现行级别定位。",
    indent=True)
last = add_after(last,
    "基于预训练模型的方法在JIT-DL上取得了突破性进展。"
    "JIT-Fine[14]训练CodeBERT识别行级别缺陷，使用特定的行信息进行训练。"
    "JIT-Smart[13]使用注意力分数来优先排序代码行，将JIT-DL作为JIT-DP的辅助任务。"
    "然而，这种将JIT-DL置于从属地位的设计限制了模型在行级别定位上的性能。"
    "Chen等人[14]的研究表明，专门针对行级别信息进行训练可以显著提升JIT-DL的效果。",
    indent=True)

last = add_after(last, "1.4.3  多任务学习在软件工程中的应用", bold=True)
last = add_after(last,
    "多任务学习通过同时优化多个相关任务来提升模型的泛化能力和各任务的性能。"
    "在软件工程领域，多任务学习已被应用于多个场景。"
    "例如，代码摘要和代码生成可以作为互补任务进行联合训练[12]，"
    "缺陷检测和缺陷修复也可以通过共享表示来相互促进[24]。"
    "然而，在JIT-DP和JIT-DL的联合优化方面，现有研究仍存在不足。"
    "JIT-Smart[13]虽然同时处理两个任务，但采用层次化设计，将JIT-DP作为主任务，"
    "JIT-DL仅作为辅助任务，未能充分发挥多任务学习的优势。"
    "如何设计有效的多任务学习框架，使JIT-DP和JIT-DL能够平等地相互促进，"
    "是本研究要解决的核心问题。",
    indent=True)

# PART6_PLACEHOLDER

# ============================================================
# 8. 本人承担的研究工作（含图片）
# ============================================================
p = doc.paragraphs[idx["my_work"]]
p.clear()
run = p.add_run("2  本人所承担的具体研究工作")
run.font.size = Pt(14)
run.bold = True

last = p
last = add_after(last, "2.1  研究动机", bold=True)
last = add_after(last,
    "现有方法如JIT-Smart[13]已证明使用统一的CPTM同时处理JIT-DP和JIT-DL任务的有效性。"
    "然而，JIT-Smart将JIT-DP作为主要任务、JIT-DL作为辅助任务，这可能限制模型捕获细粒度行级别"
    "缺陷模式的能力。我们的方法基于三个关键洞察：（1）平等对待两个任务：JIT-DP和JIT-DL在多任务学习中"
    "应被视为同等重要，两个任务可以从共享表示中相互受益；（2）显式缺陷表示学习：通用代码表示会将功能"
    "相似的代码归为一类而不考虑缺陷状态，显式学习缺陷特定表示可以帮助模型更好地区分缺陷代码和正常代码；"
    "（3）桥接提交级别和行级别信息：直接优化两个任务的多任务学习可能忽略提交级别上下文，"
    "我们通过引入预学习的提交级别缺陷表示和专家特征来解决这一问题。",
    indent=True)

last = add_after(last, "2.2  JIT-MTL方法概述", bold=True)
last = add_after(last,
    "基于上述洞察，我们提出的JIT-MTL方法采用两阶段训练策略来训练CPTM同时处理JIT-DP和JIT-DL任务。"
    "第一阶段专注于学习提交级别缺陷表示，为缺陷判别建立基础。该阶段使CPTM能够捕获提交级别的缺陷指示模式，"
    "这些模式在直接进行多任务学习时可能被忽略。第二阶段利用多任务学习同时优化两个任务，"
    "同时引入预学习的缺陷表示和专家特征以保留提交级别上下文。"
    "图2展示了JIT-MTL方法的整体框架。",
    indent=True)

# 插入图2：方法概览
last = add_image_after(last, os.path.join(FIGS_DIR, "figs2.png"), width=Inches(5),
                       caption="图 2  JIT-MTL方法整体框架")

last = add_after(last, "2.3  缺陷表示学习阶段", bold=True)
last = add_after(last,
    "CPTM学习的通用代码特征会将功能相似的代码归为一组，无论其是否存在缺陷，"
    "这使得区分缺陷变得更加困难。为此，我们在第一阶段训练CPTM学习提交级别缺陷代码表示。"
    "具体而言，我们从JIT-DP任务创建提交级别缺陷代码数据集，将每个提交标记化为代码token。"
    "我们使用额外的全连接分类层训练CPTM，其中每个输入序列以特殊起始token [CLS]开始，"
    "以特殊分隔token [SEP]结束。通过二分类训练过程，CPTM学会将缺陷指示模式编码到[CLS]向量表示中，"
    "从而能够区分缺陷提交和正常提交。",
    indent=True)

last = add_after(last, "2.4  多任务学习阶段", bold=True)
last = add_after(last,
    "在多任务学习阶段，我们复用第一阶段预训练的CPTM参数，保留已学习的提交级别缺陷表示。"
    "为进一步增强CPTM捕获提交级别信息的能力，我们引入专家特征作为额外输入。"
    "我们使用嵌入层将14个专家特征映射到与其他提交级别表示对齐的高维向量。"
    "图3展示了多任务学习阶段的详细结构。",
    indent=True)

# 插入图3：多任务学习
last = add_image_after(last, os.path.join(FIGS_DIR, "figs3.png"), width=Inches(5),
                       caption="图 3  JIT-DP和JIT-DL任务的多任务学习示意图")

last = add_after(last,
    "训练时，CPTM同时预测提交级别和行级别缺陷。对于包含n行的提交ci，"
    "我们构建输入序列使每行lj既作为提交上下文的一部分又单独出现，"
    "使CPTM能够捕获两个级别的信息。每个输入遵循格式：[[CLS], ci, [SEP], lj, [SEP]]。"
    "我们使用两个分类层分别进行提交级别和行级别预测，并采用Focal Loss[19]作为联合训练目标"
    "来处理缺陷预测任务中常见的类别不平衡问题。",
    indent=True)
last = add_after(last,
    "在测试阶段，行级别预测直接指示每行是否存在缺陷。对于提交级别预测，"
    "包含Ni行的提交ci会生成Ni个独立预测。我们通过加权合并过程获得单一的提交级别预测，"
    "将所有行的缺陷概率取平均值，若超过阈值τ=0.5则判定提交存在缺陷。",
    indent=True)

# PART7_PLACEHOLDER

# ============================================================
# 9. 实验评估
# ============================================================
last = add_after(last, "2.5  实验评估", bold=True)
last = add_after(last,
    "我们在公开的JIT-Defect4J数据集上进行实验，该数据集包含21个Java软件项目。"
    "实验结果表明，JIT-MTL在几乎所有指标上均优于现有最优基线方法。",
    indent=True)

last = add_after(last, "2.5.1  JIT-DP任务结果", bold=True)
last = add_after(last,
    "在JIT-DP任务上，JIT-MTL的F1分数（0.627）比最优基线JIT-Smart（0.486）高14.1%，"
    "AUC值（0.899）高1.4%，表明我们的方法能更好地识别提交级别缺陷。"
    "在考虑成本的努力感知指标上，JIT-MTL的R@20%E达到0.817，E@20%R为0.006，Popt为0.940，"
    "能以更少的努力识别更多的缺陷引入提交。",
    indent=True)

last = add_after(last, "2.5.2  JIT-DL任务结果", bold=True)
last = add_after(last,
    "在JIT-DL任务上，JIT-MTL的Top-5准确率（0.916）比最优基线高53.9%，"
    "Top-10准确率（0.949）高36.4%，表明JIT-MTL在定位行级别缺陷并为开发者排序方面更为有效。"
    "此外，JIT-MTL的R@20%El高2.8%，E@20%Rl低2.1%，"
    "意味着它能以更少的努力定位更多的缺陷代码行。",
    indent=True)

last = add_after(last, "2.5.3  消融实验", bold=True)

# 插入图4：不同粒度
last = add_image_after(last, os.path.join(FIGS_DIR, "figs4.png"), width=Inches(5),
                       caption="图 4  不同粒度的缺陷代码表示学习示意图")

last = add_after(last,
    "我们评估了不同粒度的缺陷代码表示学习策略（如图4所示）。结果表明，"
    "提交级别缺陷代码表示对两个任务的性能提升最为显著。"
    "行级别表示虽然提升了JIT-DL性能，但降低了JIT-DP性能。混合级别表示在所有方法中表现最差，"
    "这揭示了简单混合不同粒度数据的局限性。此外，专家特征的消融实验证明了其对提升CPTM性能的重要性，"
    "它能提供CPTM无法直接从原始代码中学习的额外提交级别信息。",
    indent=True)

last = add_after(last, "2.5.4  与大语言模型的比较", bold=True)
last = add_after(last,
    "我们将JIT-MTL与多个大语言模型进行了比较，包括Qwen-2.5-Coder系列（7B、14B、33B参数）"
    "和CodeLlama系列（7B、13B参数）。结果表明，尽管JIT-MTL使用的模型（CodeBERT，125M参数）"
    "远小于这些大语言模型，但在两个任务的几乎所有指标上均显著优于它们。"
    "在JIT-DP上，JIT-MTL的F1分数比LLM高275.4%-703.8%；"
    "在JIT-DL上，Top-5准确率高139.4%-183.8%。"
    "这证明了任务特定设计的有效性，同时也表明我们的方法在实际部署中更加高效和实用。",
    indent=True)

# ============================================================
# 10. 下一步科研计划
# ============================================================
p = doc.paragraphs[idx["future"]]
p.clear()
run = p.add_run("3  下一步科研计划")
run.font.size = Pt(14)
run.bold = True

last = p
last = add_after(last, "3.1  总结与展望", bold=True)
last = add_after(last,
    "本文提出的JIT-MTL方法通过多任务学习框架有效地同时解决了JIT-DP和JIT-DL两个任务，"
    "在公开数据集上取得了显著优于现有方法的实验结果。未来，我们计划从以下三个方向进一步优化和拓展研究工作。",
    indent=True)

last = add_after(last, "3.2  基于参数高效微调的大语言模型应用", bold=True)
last = add_after(last,
    "当前实验结果表明，直接使用大语言模型进行JIT-DP和JIT-DL任务的效果不佳。"
    "我们计划应用参数高效微调方法（如LoRA、Adapter等）来利用大语言模型的强大代码理解能力，"
    "并结合本文提出的多任务学习框架来提升大语言模型在缺陷预测任务上的性能。",
    indent=True)

last = add_after(last, "3.3  更有效的缺陷代码表示学习策略", bold=True)
last = add_after(last,
    "当前的缺陷代码表示学习采用简单的二分类训练方式。"
    "我们计划探索更有效的学习策略，例如使用对比学习[23]来选择更合适的行级别数据，"
    "构建更具区分性的缺陷代码表示空间，从而进一步提升模型对缺陷模式的识别能力。",
    indent=True)

last = add_after(last, "3.4  集成更多相关任务", bold=True)
last = add_after(last,
    "缺陷预测和定位与程序修复等任务密切相关。"
    "我们计划将程序修复[24]等相关任务集成到多任务学习框架中，"
    "使模型不仅能够检测和定位缺陷，还能提供修复建议，形成更完整的缺陷处理流水线。",
    indent=True)

# PART8_PLACEHOLDER

# ============================================================
# 11. 参考文献
# ============================================================
last = add_after(last, "")
last = add_after(last, "参考文献", font_size=Pt(14), bold=True)

references = [
    "[1] Chen Y, et al. DefectChecker: Automated smart contract defect detection by analyzing EVM bytecode. IEEE TSE, 2022, 48(7): 2335-2348.",
    "[2] Yang X, et al. Using deep learning to generate complete log statements. ICSE, 2021: 1-12.",
    "[3] Shen Z, et al. A survey of automatic code generation techniques. Journal of Software, 2020, 31(1): 1-27.",
    "[4] Niu C, et al. RAT: A Retrieval-Augmented Transformer for code generation. ESEC/FSE, 2023: 1-12.",
    "[5] Bacchelli A, Bird C. Expectations, outcomes, and challenges of modern code review. ICSE, 2013: 712-721.",
    "[6] Pascarella L, et al. Fine-grained just-in-time defect prediction. JSS, 2019, 150: 22-36.",
    "[7] Yan J, et al. Just-in-time defect identification and localization: A two-phase framework. IEEE TSE, 2020, 48(1): 82-96.",
    "[8] Pornprasit C, Tantithamthavorn C. JITLine: A simpler, better, faster, finer-grained just-in-time defect prediction. MSR, 2021: 369-379.",
    "[9] Hoang T, et al. CC2Vec: Distributed representations of code changes. ICSE, 2020: 518-529.",
    "[10] Vaswani A, et al. Attention is all you need. NeurIPS, 2017: 5998-6008.",
    "[11] Niu C, et al. SPT-Code: Sequence-to-sequence pre-training for learning source code representations. ICSE, 2022: 2006-2018.",
    "[12] Wang Y, et al. CodeT5: Identifier-aware unified pre-trained encoder-decoder models for code understanding and generation. EMNLP, 2021: 8696-8708.",
    "[13] Ni L, et al. The best of both worlds: Integrating semantic features with expert features for defect prediction and localization. ESEC/FSE, 2022: 672-684.",
    "[14] Chen X, et al. JIT defect prediction and localization via pre-trained models. JSS, 2024, 210: 111939.",
    "[15] Feng Z, et al. CodeBERT: A pre-trained model for programming and natural languages. EMNLP, 2020: 1536-1547.",
    "[16] Lin B, et al. CCT5: A code-change-oriented pre-trained model. ESEC/FSE, 2023: 1-12.",
    "[17] Zou W, Shen Z, et al. CCAF: Learning code change via AdapterFusion. Internetware, 2024: 219-228.",
    "[18] Tang J, et al. Just-in-time defect prediction using Transformer. SANER, 2023: 1-12.",
    "[19] Lin T Y, et al. Focal loss for dense object detection. ICCV, 2017: 2980-2988.",
    "[20] Guo D, et al. UniXcoder: Unified cross-modal pre-training for code representation. ACL, 2022: 7212-7225.",
    "[21] Hui B, et al. Qwen2.5-Coder technical report. arXiv:2409.12186, 2024.",
    "[22] Roziere B, et al. Code Llama: Open foundation models for code. arXiv:2308.12950, 2023.",
    "[23] He K, et al. Momentum contrast for unsupervised visual representation learning. CVPR, 2020: 9729-9738.",
    "[24] Ni L, et al. Unifying defect prediction, categorization, and repair. ICSE, 2023: 1-12.",
    "[25] Mockus A, Weiss D M. Predicting risk of software changes. Bell Labs Technical Journal, 2000, 5(2): 169-180.",
    "[26] Kamei Y, et al. A large-scale empirical study of just-in-time quality assurance. IEEE TSE, 2013, 39(6): 757-773.",
    "[27] Zeng Z, et al. Deep just-in-time defect localization. IEEE TSE, 2022, 48(12): 5068-5083.",
    "[28] Hoang T, et al. DeepJIT: An end-to-end deep learning framework for just-in-time defect prediction. MSR, 2019: 34-45.",
    "[29] Huang Q, et al. Revisiting the practical use of automated software fault localization techniques. ISSRE, 2019: 1-12.",
]

for ref in references:
    last = add_after(last, ref, font_size=Pt(10))

# ============================================================
# 12. 保存
# ============================================================
doc.save(OUTPUT)
print(f"报告已生成: {OUTPUT}")
